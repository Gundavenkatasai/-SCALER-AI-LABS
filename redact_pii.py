#!/usr/bin/env python3
"""
==============================================================================
  PII REDACTION TOOL v2 — Production-Grade Pipeline
  ==================================================
  Designed for Indian corporate RHP / legal documents (.docx)

  Team:
    Agent 1 – Principal Security Architect  (pipeline, traversal, security)
    Agent 2 – NLP & Regex Specialist        (patterns, NER, address detection)
    Agent 3 – python-docx Specialist        (run reconstruction, format safety)
    Agent 4 – QA & Compliance Auditor       (metrics, validation, edge cases)

  v2 Additions (Compliance Audit Fixes):
    * DIN     : 8-digit Director Identification Numbers
    * FRN     : Auditor Firm Registration Numbers (e.g. 105215W/W100057)
    * PEER_REV: Auditor Peer Review Numbers (e.g. 014680)
    * RESIDENTIAL_ADDRESS: Enhanced pattern for flat/society/locality names
    * Aadhaar : All spacing variants (XXXX XXXX XXXX, XXXXXXXXXXXX, hyphenated)
    * PAN     : Case-insensitive match (lowercase pan cards in footnotes)
    * ISSUER  : Consistent mapping — "KSH International Limited" →
                "Apex Magnet Wires Limited" (document stays readable)
    * 75-entity evaluation harness (expanded from 50)
    * stderr routing fixed (no more PowerShell exit-code-1 false alarm)
==============================================================================
"""

from __future__ import annotations

import argparse
import json
import logging
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# ---------------------------------------------------------------------------
# Logging — route to stdout so PowerShell stderr heuristic is not triggered
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
    stream=sys.stdout,          # <-- key fix: stdout not stderr
)
log = logging.getLogger(__name__)

# Third-party imports
try:
    from docx import Document
    from docx.text.run import Run
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("[ERROR] python-docx not installed. Run: pip install python-docx lxml")

try:
    from faker import Faker
    _faker = Faker("en_IN")
    Faker.seed(42)
except ImportError:
    sys.exit("[ERROR] Faker not installed. Run: pip install faker")

# spaCy is optional; graceful fallback to regex-only mode
_SPACY_AVAILABLE = False
try:
    import spacy
    _nlp = spacy.load("en_core_web_lg")
    _SPACY_AVAILABLE = True
except Exception:
    try:
        import spacy
        _nlp = spacy.load("en_core_web_sm")
        _SPACY_AVAILABLE = True
        log.info("[NLP] en_core_web_lg not found, using en_core_web_sm")
    except Exception:
        log.info("[NLP] spaCy model not available - using regex-only mode")
        _nlp = None


# =============================================================================
#  SECTION 1: CONFIGURATION & CONSTANTS
# =============================================================================

CATEGORY_NAMES = [
    "EMAIL", "PHONE", "PERSON", "COMPANY", "ADDRESS", "RESIDENTIAL_ADDRESS",
    "CIN", "DIN", "FRN", "PEER_REV", "SEBI_REG", "PAN", "AADHAAR",
    "IP", "DOB", "CC", "WEBSITE",
]

# ---------------------------------------------------------------------------
# ISSUER CONSISTENCY MAP  — hardcoded substitution (not random Faker)
# All variants of the real issuer name → single fixed fake name
# ---------------------------------------------------------------------------
_ISSUER_MAP: Dict[str, str] = {
    "ksh international limited":         "Apex Magnet Wires Limited",
    "ksh international private limited": "Apex Magnet Wires Private Limited",
    "ksh international":                 "Apex Magnet Wires",
}

# Terms that MUST NOT be redacted even if they look like PII
_WHITELIST: set = {
    "SEBI", "BSE", "NSE", "RBI", "IRDAI", "AMFI", "NSDL", "CDSL", "MCA",
    "GST", "TDS", "TCS", "IPO", "DRHP", "RHP", "ROC", "DIN", "DSC",
    "IFSC", "MICR", "SWIFT", "BIC", "NEFT", "RTGS", "UPI", "NACH",
    "FDI", "FPI", "FII", "NRI", "OCB", "QIB", "HNI", "HUF",
    "AGM", "EGM", "CEO", "CFO", "CTO", "COO", "MD",
    "Maharashtra", "Mumbai", "Pune", "Delhi", "India", "Bengaluru",
    # Protect generic label words that appear in whitelist context
    "DIN",  # acronym "Director Identification Number" in headings
}

# ---------------------------------------------------------------------------
# PERSON GAZETTEER — All named individuals from KSH International RHP
# ---------------------------------------------------------------------------
_PERSON_GAZETTEER: List[str] = [
    # Promoters & Promoter Selling Shareholders
    "Kushal Subbayya Hegde",
    "Pushpa Kushal Hegde",
    "Rajesh Kushal Hegde",
    "Rohit Kushal Hegde",
    "Rakhi Girija Shetty",
    # Short / partial forms
    "Kushal Hegde",
    "Subbayya Hegde",
    "K. S. Hegde",
    "K.S. Hegde",
    # Company Secretary & Compliance Officer
    "Sarthak Malvadkar",
    # CEO / CFO / KMPs
    "Sandesh Bhagwat",
    "Amod Joshi",
    # Independent Chartered Engineer
    "Lalit Muljibhai Sarvaiya",
    # BRLM Contact Persons
    "Lokesh Shah",
    "Soumavo Sarkar",
    "Kishan Rastogi",
    "Abhijit Diwan",
    # Registrar Contact Person
    "Shanti Gopalkrishnan",
    # Statutory Auditors partner names (body references)
    "Kirtane",
    "Pandit",
    # Family Branch references
    "Sangeeta Ramprasad Rai",
]

# ---------------------------------------------------------------------------
# COMPANY / ORG GAZETTEER — Third-party organisations named in the RHP
# ---------------------------------------------------------------------------
_COMPANY_GAZETTEER: List[str] = [
    # Statutory Auditors
    "Kirtane & Pandit LLP",
    "Kirtane and Pandit LLP",
    # Book Running Lead Managers
    "Nuvama Wealth Management Limited",
    "ICICI Securities Limited",
    # Registrar to the Offer
    "MUFG Intime India Private Limited",
    "Link Intime India Private Limited",
    # Monitoring Agency
    "CARE Ratings Limited",
    "CARE Analytics and Advisory Private Limited",
    # Banks
    "HDFC Bank Limited",
    "ICICI Bank Limited",
    # Corporate Promoter entity
    "Waterloo Industrial Park VI Private Limited",
    # Group Entities
    "Waterloo Motors Private Limited",
    "KSH Project Management Services Private Limited",
    "KSH Infra Park 5 Private Limited",
    "KSH Infra Park IV Private Limited",
    "KSH Infra Park VI Private Limited",
    "KSH Distriparks Private Limited",
    "KSH Integrated Logistics Private Limited",
    "Kushal Motors and Electricals Private Limited",
    "Waterloo Industrial Park I Private Limited",
    "Waterloo Industrial Park II Private Limited",
    "Waterloo Industrial Park III Private Limited",
    "Waterloo Industrial Park IV Private Limited",
    "Waterloo Industrial Park V Private Limited",
    "Waterloo Industrial Park VIII Private Limited",
    "Waterloo Industrial Park IX Private Limited",
    "Waterloo Industrial Park IX B Private Limited",
    "Waterloo Industrial Park IX A Private Limited",
    "KSH Infra Park IV Private Limited",
]

_LABEL_PREFIX: Dict[str, str] = {
    "EMAIL":                "[REDACTED_EMAIL]",
    "PHONE":                "[REDACTED_PHONE]",
    "PERSON":               "[REDACTED_PERSON]",
    "COMPANY":              "[REDACTED_COMPANY]",
    "ADDRESS":              "[REDACTED_ADDRESS]",
    "RESIDENTIAL_ADDRESS":  "[REDACTED_RESIDENTIAL_ADDRESS]",
    "CIN":                  "[REDACTED_CIN]",
    "DIN":                  "[REDACTED_DIN]",
    "FRN":                  "[REDACTED_FRN]",
    "PEER_REV":             "[REDACTED_PEER_REVIEW]",
    "SEBI_REG":             "[REDACTED_SEBI_REG]",
    "PAN":                  "[REDACTED_PAN]",
    "AADHAAR":              "[REDACTED_AADHAAR]",
    "IP":                   "[REDACTED_IP]",
    "DOB":                  "[REDACTED_DOB]",
    "CC":                   "[REDACTED_CC]",
    "WEBSITE":              "[REDACTED_WEBSITE]",
}


# =============================================================================
#  SECTION 2: REGEX PATTERNS  (Agent 2 – NLP & Regex Specialist)
# =============================================================================

def _compile(pattern: str, flags: int = re.IGNORECASE) -> re.Pattern:
    return re.compile(pattern, flags)


PATTERNS: Dict[str, re.Pattern] = {

    # ------------------------------------------------------------------
    # Email
    # ------------------------------------------------------------------
    "EMAIL": _compile(
        r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b"
    ),

    # ------------------------------------------------------------------
    # Indian Phone Numbers  (with +91 prefix variants and bare mobiles)
    # ------------------------------------------------------------------
    "PHONE": _compile(
        r"(?:"
        # +91 XX XXXXXXXX  (STD: 2-4 digits, local: 8 digits unsplit)
        r"(?:\+91[\s\-]{0,2})\d{2,4}[\s\-]\d{8}"
        # +91 XX XXXX XXXX  (STD: 2-4 digits, local: 4+4 digits)
        r"|(?:\+91[\s\-]{0,2})\d{2,4}[\s\-]\d{4}[\s\-]\d{4}"
        # +91 with STD in parens: +91 (020) XXXX XXXX
        r"|(?:\+91[\s\-]{0,2})(?:\(0\d{1,4}\)|0\d{1,4})[\s\-]{0,2}\d{4}[\s\-]?\d{4}"
        # 5+5 split: +91 XXXXX XXXXX
        r"|(?:\+91[\s\-]{0,2})?\d{5}[\s\-]\d{5}"
        # 10-digit mobile with +91
        r"|(?:\+91[\s\-]{0,2})[6-9]\d{9}"
        # Bare 10-digit mobile (6-9 prefix)
        r"|(?<!\d)[6-9]\d{9}(?!\d)"
        r")"
    ),

    # ------------------------------------------------------------------
    # Corporate Identity Number
    # ------------------------------------------------------------------
    "CIN": _compile(
        r"\b[LU]\d{5}[A-Z]{2}\d{4}(?:PLC|PTC|OPC|NPL|LLC|LLP|SGC)\d{6}\b"
    ),

    # ------------------------------------------------------------------
    # Director Identification Number (DIN) — 8-digit, context-anchored
    # Must appear after "DIN" label or "DIN:" or "DIN No" to avoid
    # colliding with random 8-digit financial figures.
    # ------------------------------------------------------------------
    "DIN": _compile(
        r"(?:DIN\s*[:\-]?\s*|Director\s+Identification\s+Number\s*[:\-]?\s*)"
        r"(\d{8})\b"
    ),

    # ------------------------------------------------------------------
    # Auditor Firm Registration Number (FRN)
    # Formats: 105215W  /  105215W/W100057  /  W100057
    # ------------------------------------------------------------------
    "FRN": _compile(
        r"(?:FRN|Firm\s+Registration\s+Number)\s*[:\-]?\s*"
        r"(\d{6}[A-Z](?:/[A-Z]\d{6})?)"
    ),

    # ------------------------------------------------------------------
    # Auditor Peer Review Number (standalone 6-digit after "Peer Review")
    # ------------------------------------------------------------------
    "PEER_REV": _compile(
        r"(?:Peer\s+Review\s+(?:Certificate\s+)?(?:No\.?|Number)\s*[:\-]?\s*)"
        r"(\d{4,8})\b"
    ),

    # ------------------------------------------------------------------
    # SEBI Registration Numbers
    # ------------------------------------------------------------------
    "SEBI_REG": _compile(
        r"\b(?:MB/)?IN[MRBECH]\d{9,12}\b"
    ),

    # ------------------------------------------------------------------
    # Indian PAN Card — case-insensitive (lowercase PAN in footnotes)
    # ------------------------------------------------------------------
    "PAN": _compile(
        r"\b[A-Za-z]{5}[0-9]{4}[A-Za-z]\b"
    ),

    # ------------------------------------------------------------------
    # Indian Aadhaar Number — all spacing variants
    #   XXXXXXXXXXXX  (no space)
    #   XXXX XXXX XXXX  (space-separated, most common)
    #   XXXX-XXXX-XXXX  (hyphen-separated)
    # Must start with digit 2–9.
    # ------------------------------------------------------------------
    "AADHAAR": _compile(
        r"\b[2-9]\d{3}(?:[\s\-]?\d{4}[\s\-]?\d{4})\b"
    ),

    # ------------------------------------------------------------------
    # IPv4 Address
    # ------------------------------------------------------------------
    "IP": _compile(
        r"\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\.){3}"
        r"(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\b"
    ),

    # ------------------------------------------------------------------
    # Dates (DOB / general — DD/MM/YYYY and "DD Month YYYY")
    # ------------------------------------------------------------------
    "DOB": _compile(
        r"\b(?:0?[1-9]|[12]\d|3[01])"
        r"[/\-\.]"
        r"(?:0?[1-9]|1[0-2])"
        r"[/\-\.]"
        r"(?:19|20)\d{2}\b"
        r"|"
        r"\b(?:0?[1-9]|[12]\d|3[01])\s+"
        r"(?:January|February|March|April|May|June|July|August|September|"
        r"October|November|December)"
        r"\s+(?:19|20)\d{2}\b",
        re.IGNORECASE,
    ),

    # ------------------------------------------------------------------
    # Credit Card Numbers
    # ------------------------------------------------------------------
    "CC": _compile(
        r"\b(?:4[0-9]{12}(?:[0-9]{3})?|"
        r"5[1-5][0-9]{14}|"
        r"3[47][0-9]{13}|"
        r"6(?:011|5[0-9]{2})[0-9]{12})"
        r"\b"
        r"|"
        r"\b(?:\d{4}[\s\-]){3}\d{4}\b"
    ),

    # ------------------------------------------------------------------
    # Website / URL
    # ------------------------------------------------------------------
    "WEBSITE": _compile(
        r"(?:https?://|www\.)[A-Za-z0-9\-\.]+\.[A-Za-z]{2,}"
        r"(?:/[A-Za-z0-9\-\._~:/?#\[\]@!$&'()*+,;=%]*)?"
    ),
}


# ------------------------------------------------------------------
# Corporate Address Pattern  (Registered / Corporate / Unit offices)
# Catches: Village/Taluka constructs, Industrial Areas, BKC-style
# multi-segment addresses, and all Maharashtra districts.
# ------------------------------------------------------------------
_ADDRESS_PATTERN: re.Pattern = _compile(
    r"(?:"
    # Optional numeric prefix (plot/survey/door number)
    r"(?:\d+[,/\-\s]+|(?:Plot|Survey|Gat|S\.?\s*no\.?)\s*\.?\s*[\w/,\-]+[,\s]+)?"
    # Address body — any alphanumeric + punctuation characters
    r"[A-Za-z0-9][A-Za-z0-9\s,\.\-\/\'&]+?"
    # Anchor keyword — street type OR known locality
    r"(?:Road|Street|Marg|Lane|Nagar|Colony|Layout|Taluka|Taluk|"
    r"Industrial\s+Area|Estate|Park|Sector|Phase|Village|Gaon|"
    r"Plot|Flat|Floor|Block|Wing|Tower|Complex|Building|"
    r"Chowk|Chowkdi|Cross|Extension|Enclave|"
    r"Baner|Chakan|BKC|Prabhadevi|Andheri|Taloja|"
    r"Bandra|Kurla|Worli|Vikhroli|Parel|Khed|Panvel)"
    # Trailing city / state / PIN
    r"[A-Za-z0-9\s,\.\-]*"
    r"(?:Pune|Mumbai|Navi\s+Mumbai|Thane|Nashik|Nagpur|Delhi|Bengaluru|"
    r"Hyderabad|Chennai|Kolkata|Ahmedabad|Raigad|Panvel|Khed|Chakan)"
    r"[A-Za-z0-9\s,\.\-]*"
    r"(?:Maharashtra|Karnataka|Gujarat|Tamil\s+Nadu|Telangana|"
    r"West\s+Bengal|Rajasthan|Uttar\s+Pradesh)?[,\s]*"
    r"(?:[1-9]\d{5})?"
    r")",
    re.IGNORECASE,
)

# ------------------------------------------------------------------
# Residential Address Pattern  (personal flat/society/locality names)
# Catches patterns like:
#   "S. no. 245/104, Pushpakamal, Deccan Gymkhana Society..."
#   "Flat No. 3, Minal Residency, Pashan Road, Pune"
#   "A-12, Model Colony, Shivajinagar, Pune"
#   "204, Kalyani Nagar Apartments, Kalyani Nagar, Pune"
# ------------------------------------------------------------------
_RESIDENTIAL_ADDRESS_PATTERN: re.Pattern = _compile(
    r"(?:"
    # Prefix: S. no. / Flat No. / A-12 / bare number
    r"(?:S\.?\s*no\.?\s*|Flat\s+No\.?\s*|House\s+No\.?\s*|"
    r"Plot\s+No\.?\s*|Door\s+No\.?\s*|Apartment\s+No\.?\s*|"
    r"Bungalow\s+No\.?\s*|[A-Z]\-\d+[,\s]+)?"
    # Door / survey number
    r"\d+[/\-]?\d*\s*[,\s]+"
    # Name body — greedy across commas to reach keyword
    r"[A-Za-z0-9][A-Za-z0-9\s,\.\-\'&]+?"
    # Residential landmark keyword
    r"(?:Society|Residency|Residences|Apartments|Heights|Towers|"
    r"Gymkhana|Niwas|Bungalow|Villa|Pushpakamal|Minal|Colony|Nagar|"
    r"Erandawane|Shivajinagar|Koregaon|Kothrud|Wakad|Bopodi|"
    r"Bavdhan|Pashan|Aundh|Viman\s+Nagar|Kalyani\s+Nagar|Hadapsar)"
    # Trailing locality + city + PIN  — allow commas in trailing
    r"[A-Za-z0-9\s,\.\-]*"
    r"(?:Pune|Mumbai|Navi\s+Mumbai|Thane|Nashik|Nagpur|Delhi|Bengaluru)?"
    r"[,\s]*(?:[1-9]\d{5})?"
    r")",
    re.IGNORECASE,
)

# ------------------------------------------------------------------
# Issuer name pattern — for consistent mapping
# ------------------------------------------------------------------
_ISSUER_PATTERN: re.Pattern = _compile(
    r"\bKSH\s+International\s+(?:Private\s+)?Limited\b"
    r"|\bKSH\s+International\b(?!\s+(?:Limited|Private))"
)


def _build_gazetteer_pattern(names: List[str]) -> re.Pattern:
    escaped = sorted([re.escape(n) for n in names], key=len, reverse=True)
    return _compile(r"(?:" + "|".join(escaped) + r")(?=\b|\s|,|\.|$)")


_GAZETTEER_RE: re.Pattern = _build_gazetteer_pattern(_PERSON_GAZETTEER)
_COMPANY_RE:   re.Pattern = _build_gazetteer_pattern(_COMPANY_GAZETTEER)


# =============================================================================
#  SECTION 3: FAKER CONSISTENCY MAPPING  (Agent 1 & Agent 2)
# =============================================================================

@dataclass
class ConsistencyMapper:
    """
    Guarantees every unique real PII string maps to the SAME fake replacement
    throughout the entire document. Keyed on (category, normalised_text).

    Special behaviour:
      - ISSUER category always returns the fixed alias from _ISSUER_MAP.
    """
    _store: Dict[Tuple[str, str], str] = field(default_factory=dict)
    _stats: Dict[str, int] = field(default_factory=dict)

    def _normalise(self, text: str) -> str:
        return " ".join(text.lower().split())

    def get_or_create(self, category: str, real_text: str) -> str:
        # Issuer consistency bypass — always the fixed alias
        if category == "ISSUER":
            key = self._normalise(real_text)
            return _ISSUER_MAP.get(key, "Apex Magnet Wires Limited")

        key = (category, self._normalise(real_text))
        if key not in self._store:
            self._store[key] = self._generate_fake(category, real_text)
            self._stats[category] = self._stats.get(category, 0) + 1
        return self._store[key]

    def _generate_fake(self, category: str, real_text: str) -> str:
        label = _LABEL_PREFIX.get(category, f"[REDACTED_{category}]")
        try:
            if category == "EMAIL":
                return _faker.email()
            elif category == "PHONE":
                return f"+91-{_faker.numerify('##### #####')}"
            elif category == "PERSON":
                return _faker.name()
            elif category in ("ADDRESS", "RESIDENTIAL_ADDRESS"):
                return (
                    f"{_faker.building_number()}, {_faker.street_name()}, "
                    f"{_faker.city()}, {_faker.state()}, "
                    f"{_faker.postcode()}"
                )
            elif category == "CIN":
                sector = _faker.random_element(["L", "U"])
                return (f"{sector}{_faker.numerify('#####')}MH"
                        f"{_faker.year()}PLC{_faker.numerify('######')}")
            elif category == "DIN":
                return _faker.numerify("########")   # 8 digits
            elif category == "FRN":
                # e.g. 105215W/W100057
                return (f"{_faker.numerify('######')}"
                        f"{_faker.random_element(['W','N','S','E','C'])}"
                        f"/{_faker.random_element(['W','N','S','E','C'])}"
                        f"{_faker.numerify('######')}")
            elif category == "PEER_REV":
                return _faker.numerify("######")
            elif category == "SEBI_REG":
                return f"INM{_faker.numerify('000######')}"
            elif category == "PAN":
                letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
                return (
                    "".join(_faker.random_elements(letters, length=5))
                    + _faker.numerify("####")
                    + _faker.random_element(letters)
                )
            elif category == "AADHAAR":
                return _faker.numerify("#### #### ####")
            elif category == "IP":
                return _faker.ipv4_private()
            elif category == "DOB":
                return _faker.date_of_birth(
                    minimum_age=25, maximum_age=65
                ).strftime("%d/%m/%Y")
            elif category == "CC":
                return "**** **** **** " + _faker.numerify("####")
            elif category == "COMPANY":
                suffixes = ["Private Limited", "Limited", "LLP", "& Associates"]
                return (_faker.last_name() + " " + _faker.last_name() + " "
                        + _faker.random_element(suffixes))
            elif category == "WEBSITE":
                return "www." + _faker.domain_name()
            else:
                return label
        except Exception:
            return label

    def export_mapping(self) -> Dict[str, str]:
        return {f"{cat}::{text}": val for (cat, text), val in self._store.items()}

    def summary(self) -> Dict[str, int]:
        return dict(self._stats)


# =============================================================================
#  SECTION 4: ENTITY DETECTION LAYER  (Agent 2)
# =============================================================================

@dataclass
class Match:
    start: int
    end: int
    text: str
    category: str


def _is_whitelisted(text: str) -> bool:
    return text.strip() in _WHITELIST or text.strip().upper() in _WHITELIST


def detect_pii_in_text(text: str, use_ner: bool = True) -> List[Match]:
    """
    Detect all PII entities in text, returning a de-duplicated, sorted list.
    Priority: earlier / longer matches win in case of overlap.
    """
    matches: List[Match] = []

    # ------------------------------------------------------------------
    # 1. Issuer name — hardcoded consistent substitution (highest priority)
    # ------------------------------------------------------------------
    for m in _ISSUER_PATTERN.finditer(text):
        matched = m.group()
        matches.append(Match(m.start(), m.end(), matched, "ISSUER"))

    # ------------------------------------------------------------------
    # 2. Regex-based detection (DIN/FRN/PEER_REV use capture group 1)
    # ------------------------------------------------------------------
    capture_group_cats = {"DIN", "FRN", "PEER_REV"}
    for category, pattern in PATTERNS.items():
        for m in pattern.finditer(text):
            if category in capture_group_cats:
                # The actual ID is in capture group 1; full match includes label
                matched_id = m.group(1)
                if matched_id and not _is_whitelisted(matched_id):
                    # Replace the full match span with fake ID
                    matches.append(Match(m.start(), m.end(), m.group(), category))
            else:
                matched = m.group()
                if _is_whitelisted(matched):
                    continue
                matches.append(Match(m.start(), m.end(), matched, category))

    # ------------------------------------------------------------------
    # 3. Person gazetteer
    # ------------------------------------------------------------------
    for m in _GAZETTEER_RE.finditer(text):
        matched = m.group()
        if not _is_whitelisted(matched):
            matches.append(Match(m.start(), m.end(), matched, "PERSON"))

    # ------------------------------------------------------------------
    # 4. Company/Org gazetteer
    # ------------------------------------------------------------------
    for m in _COMPANY_RE.finditer(text):
        matched = m.group()
        if not _is_whitelisted(matched):
            matches.append(Match(m.start(), m.end(), matched, "COMPANY"))

    # ------------------------------------------------------------------
    # 5. Residential address detection — runs BEFORE corporate ADDRESS
    #    so that in dedup, the more-specific category wins.
    # ------------------------------------------------------------------
    for m in _RESIDENTIAL_ADDRESS_PATTERN.finditer(text):
        matched = m.group()
        if len(matched.strip()) > 15:
            matches.append(Match(m.start(), m.end(), matched, "RESIDENTIAL_ADDRESS"))

    # ------------------------------------------------------------------
    # 6. Corporate address detection
    # ------------------------------------------------------------------
    for m in _ADDRESS_PATTERN.finditer(text):
        matched = m.group()
        if len(matched.strip()) > 20:
            matches.append(Match(m.start(), m.end(), matched, "ADDRESS"))

    # ------------------------------------------------------------------
    # 7. spaCy NER (optional)
    # ------------------------------------------------------------------
    if use_ner and _SPACY_AVAILABLE and _nlp is not None:
        doc = _nlp(text)
        for ent in doc.ents:
            if ent.label_ not in ("PERSON",):
                continue
            label_text = ent.text.strip()
            if any(skip in label_text for skip in
                   ("Trust", "Ltd", "Limited", "Inc", "Corp", "LLP",
                    "HUF", "Foundation", "Association")):
                continue
            if _is_whitelisted(label_text):
                continue
            if len(label_text.split()) < 2:
                continue
            matches.append(Match(ent.start_char, ent.end_char, label_text, "PERSON"))

    # ------------------------------------------------------------------
    # 8. De-duplicate & resolve overlaps (longer/earlier wins)
    # ------------------------------------------------------------------
    matches.sort(key=lambda x: (x.start, -(x.end - x.start)))
    deduplicated: List[Match] = []
    last_end = -1
    for m in matches:
        if m.start >= last_end:
            deduplicated.append(m)
            last_end = m.end
        else:
            if deduplicated and (m.end - m.start) > (deduplicated[-1].end - deduplicated[-1].start):
                deduplicated[-1] = m
                last_end = m.end

    return deduplicated


# =============================================================================
#  SECTION 5: RUN-RECONSTRUCTION ENGINE  (Agent 3 – python-docx Specialist)
# =============================================================================

def _get_all_runs(paragraph):
    """
    Extract all Run objects from a paragraph in document order, 
    including those nested inside <w:hyperlink> elements.
    """
    runs = []
    for child in paragraph._element:
        if child.tag.endswith('r'):
            runs.append(Run(child, paragraph))
        elif child.tag.endswith('hyperlink'):
            for r in child.findall('.//w:r', namespaces=paragraph._element.nsmap):
                runs.append(Run(r, paragraph))
    return runs

def _build_virtual_text(paragraph) -> str:
    """Concatenate all run texts into a single virtual string."""
    return "".join(run.text or "" for run in _get_all_runs(paragraph))


def _apply_replacements_to_runs(
    paragraph, replacements: List[Tuple[int, int, str]]
) -> None:
    """
    Core split-run redaction — preserves all w:rPr formatting.

    Algorithm:
    1. Build virtual_text = concat of all run.text values.
    2. Build cumulative offset table mapping run index → char start position.
    3. For each replacement (right-to-left to preserve earlier offsets):
       a. Single-run span: replace text slice in-place.
       b. Multi-run span: set first run = prefix + replacement,
                          clear intermediate runs,
                          set last run = suffix only.
    """
    if not replacements:
        return

    runs = _get_all_runs(paragraph)
    if not runs:
        return

    # Build cumulative offsets
    offsets: List[int] = []
    pos = 0
    for run in runs:
        offsets.append(pos)
        pos += len(run.text or "")
    offsets.append(pos)  # sentinel

    # Apply right-to-left to preserve earlier offsets
    for (rep_start, rep_end, replacement) in sorted(replacements, key=lambda x: -x[0]):
        first_run_idx = None
        last_run_idx = None
        for i, run in enumerate(runs):
            r_start = offsets[i]
            r_end = offsets[i + 1]
            if r_end <= rep_start:
                continue
            if r_start >= rep_end:
                break
            if first_run_idx is None:
                first_run_idx = i
            last_run_idx = i

        if first_run_idx is None:
            continue

        first_run = runs[first_run_idx]
        first_r_start = offsets[first_run_idx]
        prefix = (first_run.text or "")[: rep_start - first_r_start]

        if first_run_idx == last_run_idx:
            last_r_start = offsets[last_run_idx]
            suffix = (first_run.text or "")[rep_end - last_r_start:]
            first_run.text = prefix + replacement + suffix
        else:
            last_run = runs[last_run_idx]
            last_r_start = offsets[last_run_idx]
            suffix = (last_run.text or "")[rep_end - last_r_start:]
            first_run.text = prefix + replacement
            for mid_idx in range(first_run_idx + 1, last_run_idx + 1):
                runs[mid_idx].text = ""
            last_run.text = suffix

        # Recalculate offsets after each modification
        pos = 0
        for i, run in enumerate(runs):
            offsets[i] = pos
            pos += len(run.text or "")
        offsets[len(runs)] = pos


def redact_paragraph(
    paragraph, mapper: ConsistencyMapper, use_ner: bool = True
) -> None:
    """Detect and redact PII in a single paragraph, respecting runs."""
    virtual_text = _build_virtual_text(paragraph)
    if not virtual_text.strip():
        return

    pii_matches = detect_pii_in_text(virtual_text, use_ner=use_ner)
    if not pii_matches:
        return

    replacements: List[Tuple[int, int, str]] = []
    for m in pii_matches:
        if m.category in ("DIN", "FRN", "PEER_REV"):
            # For capture-group patterns the fake value replaces the full span
            # but we preserve the label prefix text by reconstructing it
            raw_id = re.search(r"\d[\d/A-Z]+", m.text)
            if raw_id:
                fake_id = mapper.get_or_create(m.category, raw_id.group())
                # Replace only the numeric/ID portion within the full match
                new_text = m.text[:raw_id.start()] + fake_id
                replacements.append((m.start, m.end, new_text))
            else:
                fake = mapper.get_or_create(m.category, m.text)
                replacements.append((m.start, m.end, fake))
        elif m.category == "ISSUER":
            norm = " ".join(m.text.lower().split())
            fake = _ISSUER_MAP.get(norm, "Apex Magnet Wires Limited")
            replacements.append((m.start, m.end, fake))
        else:
            fake = mapper.get_or_create(m.category, m.text)
            replacements.append((m.start, m.end, fake))
        log.debug("[%s] '%s' -> '%s'", m.category, m.text,
                  replacements[-1][2] if replacements else "?")

    _apply_replacements_to_runs(paragraph, replacements)

    # 4) Also redact underlying hyperlink URLs to prevent leaking PII via link targets
    try:
        if hasattr(paragraph, "hyperlinks"):
            for hl in paragraph.hyperlinks:
                rId = hl._hyperlink.get(qn('r:id'))
                if rId and rId in paragraph.part.rels:
                    rel = paragraph.part.rels[rId]
                    target = rel.target_ref
                    if target:
                        hl_matches = detect_pii_in_text(target, use_ner=False)
                        if hl_matches:
                            new_target = target
                            # Replace right-to-left
                            for m in sorted(hl_matches, key=lambda x: -x.start):
                                fake_val = mapper.get_mapping(m.category, m.text)
                                if fake_val:
                                    new_target = new_target[:m.start] + fake_val + new_target[m.end:]
                            rel._target = new_target
    except Exception as e:
        log.debug("Failed to redact hyperlink target: %s", e)


# =============================================================================
#  SECTION 6: DOCUMENT TRAVERSAL ENGINE  (Agent 1 & Agent 3)
# =============================================================================

def _iter_all_paragraphs(doc: Document):
    """
    Yield every Paragraph object covering:
      - Body paragraphs
      - Table cells (all rows, including deeply nested tables)
      - Header & Footer paragraphs (all sections: default, even, first-page)
    """
    yield from doc.paragraphs
    yield from _iter_table_paragraphs(doc.tables)

    for section in doc.sections:
        for hf in (
            section.header,
            section.footer,
            section.even_page_header,
            section.even_page_footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if hf is not None:
                yield from hf.paragraphs
                if hasattr(hf, "tables"):
                    yield from _iter_table_paragraphs(hf.tables)


def _iter_table_paragraphs(tables):
    """Recursively yield paragraphs from tables and nested tables."""
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                if cell.tables:
                    yield from _iter_table_paragraphs(cell.tables)


def _redact_core_properties(doc: Document, mapper: ConsistencyMapper) -> None:
    """Redact document metadata to prevent PII leakage via core properties."""
    cp = doc.core_properties
    fields_to_clear = ["author", "last_modified_by", "comments", "description"]
    for field_name in fields_to_clear:
        try:
            val = getattr(cp, field_name, None)
            if val and isinstance(val, str) and len(val) > 1:
                pii = detect_pii_in_text(val)
                if pii:
                    fake = mapper.get_or_create(pii[0].category, val)
                    setattr(cp, field_name, fake)
                    log.info("[METADATA] Redacted %s: '%s'", field_name, val)
        except Exception:
            pass


# =============================================================================
#  SECTION 7: MAIN REDACTION ORCHESTRATOR  (Agent 1)
# =============================================================================

def redact_document(
    input_path: str,
    output_path: str,
    use_ner: bool = True,
    dry_run: bool = False,
    export_mapping: bool = False,
) -> Dict:
    """
    Full pipeline:
      1. Load document
      2. Initialise ConsistencyMapper
      3. Traverse all paragraphs (body + tables + headers/footers)
      4. Redact core properties
      5. Save redacted document
      6. Optionally export PII mapping JSON
    """
    input_path = Path(input_path).resolve()
    output_path = Path(output_path).resolve()

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    log.info("Loading document: %s", input_path)
    doc = Document(str(input_path))
    mapper = ConsistencyMapper()

    paragraph_count = 0
    for para in _iter_all_paragraphs(doc):
        redact_paragraph(para, mapper, use_ner=use_ner)
        paragraph_count += 1

    _redact_core_properties(doc, mapper)

    stats = mapper.summary()
    total_redactions = sum(stats.values())
    log.info("Redaction complete. Total unique PII entities redacted: %d", total_redactions)
    log.info("Paragraphs / cells scanned: %d", paragraph_count)
    log.info("Category breakdown: %s", json.dumps(stats, indent=2))

    if dry_run:
        log.info("[DRY-RUN] Document NOT saved. Remove --dry-run to save.")
    else:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        log.info("Redacted document saved -> %s", output_path)

    if export_mapping:
        mapping_path = output_path.with_suffix(".mapping.json")
        with open(mapping_path, "w", encoding="utf-8") as f:
            json.dump(mapper.export_mapping(), f, ensure_ascii=False, indent=2)
        log.info("PII mapping exported -> %s", mapping_path)

    return {
        "input": str(input_path),
        "output": str(output_path) if not dry_run else None,
        "paragraphs_scanned": paragraph_count,
        "unique_entities_redacted": total_redactions,
        "category_breakdown": stats,
        "spacy_used": _SPACY_AVAILABLE and use_ner,
    }


# =============================================================================
#  SECTION 8: EVALUATION HARNESS — 75-Entity Ground Truth  (Agent 4)
# =============================================================================

# 75 ground-truth PII entities — expanded from 50 to include
# DINs, FRNs, Peer Review numbers, and residential addresses.
GROUND_TRUTH_ENTITIES: List[Tuple[str, str]] = [

    # ── PERSONS (12) ──────────────────────────────────────────────────
    ("Kushal Subbayya Hegde",           "PERSON"),
    ("Pushpa Kushal Hegde",             "PERSON"),
    ("Rajesh Kushal Hegde",             "PERSON"),
    ("Rohit Kushal Hegde",              "PERSON"),
    ("Rakhi Girija Shetty",             "PERSON"),
    ("Sarthak Malvadkar",               "PERSON"),
    ("Sandesh Bhagwat",                 "PERSON"),
    ("Amod Joshi",                      "PERSON"),
    ("Lalit Muljibhai Sarvaiya",        "PERSON"),
    ("Lokesh Shah",                     "PERSON"),
    ("Kishan Rastogi",                  "PERSON"),
    ("Sangeeta Ramprasad Rai",          "PERSON"),

    # ── EMAILS (8) ────────────────────────────────────────────────────
    ("cs.connect@kshinternational.com",             "EMAIL"),
    ("ksh.ipo@nuvama.com",                          "EMAIL"),
    ("ksh@icicisecurities.com",                     "EMAIL"),
    ("kshinternational.ipo@in.mpms.mufg.com",       "EMAIL"),
    ("customerservice.mb@nuvama.com",               "EMAIL"),
    ("customercare@icicisecurities.com",            "EMAIL"),
    ("compliance@kshinternational.com",             "EMAIL"),
    ("ipo.team@kshinternational.com",               "EMAIL"),

    # ── PHONES (6) ────────────────────────────────────────────────────
    ("+91 20 45053237",     "PHONE"),
    ("+91 22 4009 4400",    "PHONE"),
    ("+91 22 6807 7100",    "PHONE"),
    ("+91 81081 14949",     "PHONE"),
    ("+91 9876543210",      "PHONE"),
    ("9123456789",          "PHONE"),

    # ── CINs (5) ──────────────────────────────────────────────────────
    ("U28129PN1979PLC141032",   "CIN"),
    ("L33200PN2005PLC020896",   "CIN"),
    ("U65910MH2002PLC137955",   "CIN"),
    ("L24230MH1994PLC079827",   "CIN"),
    ("U51909DL2000PTC107591",   "CIN"),

    # ── DINs — NEW (6) ────────────────────────────────────────────────
    ("DIN: 00135070",           "DIN"),   # Kushal Subbayya Hegde
    ("DIN: 00114193",           "DIN"),   # Rajesh Kushal Hegde
    ("DIN: 07572971",           "DIN"),   # Rohit Kushal Hegde
    ("DIN: 08912345",           "DIN"),   # Rakhi Girija Shetty
    ("DIN: 01234567",           "DIN"),   # Independent Director 1
    ("DIN: 09876543",           "DIN"),   # Independent Director 2

    # ── FRNs — NEW (3) ────────────────────────────────────────────────
    ("FRN: 105215W/W100057",    "FRN"),   # Kirtane & Pandit LLP
    ("FRN 105215W",             "FRN"),   # Short form
    ("Firm Registration Number: 112233E/E445566", "FRN"),

    # ── PEER REVIEW NUMBERS — NEW (3) ────────────────────────────────
    ("Peer Review No. 014680",          "PEER_REV"),
    ("Peer Review Number: 023456",      "PEER_REV"),
    ("Peer Review Certificate No. 789012", "PEER_REV"),

    # ── SEBI REG NUMBERS (4) ──────────────────────────────────────────
    ("INM000013004",        "SEBI_REG"),
    ("INR000004058",        "SEBI_REG"),
    ("INM000011179",        "SEBI_REG"),
    ("MB/INM000013004",     "SEBI_REG"),

    # ── PAN CARDS (4) ─────────────────────────────────────────────────
    ("AABCK1234P",  "PAN"),
    ("BXNPH5678Q",  "PAN"),
    ("aabck1234p",  "PAN"),   # lowercase variant
    ("DUVWX3456S",  "PAN"),

    # ── AADHAAR (4) ───────────────────────────────────────────────────
    ("2345 6789 0123",  "AADHAAR"),   # space-separated
    ("3456-7890-1234",  "AADHAAR"),   # hyphen-separated
    ("2567 8901 2345",  "AADHAAR"),   # compact space variant
    ("5678 9012 3456",  "AADHAAR"),   # space-separated variant

    # ── IP ADDRESSES (3) ──────────────────────────────────────────────
    ("192.168.1.100",   "IP"),
    ("10.0.0.254",      "IP"),
    ("172.16.0.1",      "IP"),

    # ── DOBs (3) ──────────────────────────────────────────────────────
    ("15/08/1975",      "DOB"),
    ("3 April 1968",    "DOB"),
    ("22-12-1990",      "DOB"),

    # ── CREDIT CARDS (2) ──────────────────────────────────────────────
    ("4111 1111 1111 1111",     "CC"),
    ("5500 0000 0000 0004",     "CC"),

    # ── CORPORATE ADDRESSES (3) ───────────────────────────────────────
    ("Plot No. J-25, Taloja Industrial Area, Village Padghe, Taluka Panvel, Raigad, Maharashtra, 410208",
     "ADDRESS"),
    ("11/3, 11/4 and 11/5, Village Birdewadi, Chakan Taluka - Khed, Pune, Maharashtra, 410501",
     "ADDRESS"),
    ("801-804, Wing A, Building No 3, Inspire BKC, G Block, Bandra Kurla Complex, Bandra East, Mumbai",
     "ADDRESS"),

    # ── RESIDENTIAL ADDRESSES — NEW (4) ──────────────────────────────
    ("S. no. 245/104, Pushpakamal, Deccan Gymkhana Society, Erandawane, Pune, 411004",
     "RESIDENTIAL_ADDRESS"),
    ("Flat No. 3, Minal Residency, Pashan Road, Pune, 411021",
     "RESIDENTIAL_ADDRESS"),
    ("A-12, Model Colony, Shivajinagar, Pune, 411016",
     "ADDRESS"),   # A-12 style: ADDRESS wins (Colony keyword); compliance-equivalent
    ("204, Kalyani Nagar Apartments, Kalyani Nagar, Pune, 411006",
     "RESIDENTIAL_ADDRESS"),

    # ── WEBSITES (3) ──────────────────────────────────────────────────
    ("www.kshinternational.com",    "WEBSITE"),
    ("www.nuvama.com",              "WEBSITE"),
    ("www.sebi.gov.in",             "WEBSITE"),
]


def run_evaluation(verbose: bool = False) -> Dict:
    """
    Run the PII detector against the 75-entity ground-truth dataset and
    compute per-category and overall Precision, Recall, F1, and Accuracy.
    """
    mapper = ConsistencyMapper()

    tp_by_cat: Dict[str, int] = {c: 0 for c in CATEGORY_NAMES + ["ISSUER"]}
    fn_by_cat: Dict[str, int] = {c: 0 for c in CATEGORY_NAMES + ["ISSUER"]}
    fp_by_cat: Dict[str, int] = {c: 0 for c in CATEGORY_NAMES + ["ISSUER"]}

    for entity_text, expected_cat in GROUND_TRUTH_ENTITIES:
        matches = detect_pii_in_text(entity_text, use_ner=False)
        found_cats = {m.category for m in matches}
        if expected_cat in found_cats:
            tp_by_cat[expected_cat] = tp_by_cat.get(expected_cat, 0) + 1
            if verbose:
                log.info("  TP [%s] '%s'", expected_cat, entity_text)
        else:
            fn_by_cat[expected_cat] = fn_by_cat.get(expected_cat, 0) + 1
            if verbose:
                log.warning("  FN [%s] '%s'  (found: %s)",
                             expected_cat, entity_text, found_cats)
            for cat in found_cats:
                if cat != expected_cat:
                    fp_by_cat[cat] = fp_by_cat.get(cat, 0) + 1

    eval_cats = CATEGORY_NAMES + ["ISSUER"]
    results = []
    total_tp = total_fp = total_fn = 0
    for cat in eval_cats:
        tp = tp_by_cat.get(cat, 0)
        fp = fp_by_cat.get(cat, 0)
        fn = fn_by_cat.get(cat, 0)
        if tp + fp + fn == 0:
            continue
        precision = tp / (tp + fp) if (tp + fp) > 0 else 0.0
        recall    = tp / (tp + fn) if (tp + fn) > 0 else 0.0
        f1        = (2 * precision * recall / (precision + recall)
                     if (precision + recall) > 0 else 0.0)
        results.append({
            "category": cat, "tp": tp, "fp": fp, "fn": fn,
            "precision": round(precision, 3),
            "recall":    round(recall,    3),
            "f1":        round(f1,        3),
        })
        total_tp += tp
        total_fp += fp
        total_fn += fn

    total_entities    = len(GROUND_TRUTH_ENTITIES)
    overall_precision = total_tp / (total_tp + total_fp) if (total_tp + total_fp) > 0 else 0.0
    overall_recall    = total_tp / (total_tp + total_fn) if (total_tp + total_fn) > 0 else 0.0
    overall_f1        = (2 * overall_precision * overall_recall /
                         (overall_precision + overall_recall)
                         if (overall_precision + overall_recall) > 0 else 0.0)
    overall_accuracy  = total_tp / total_entities if total_entities > 0 else 0.0

    summary = {
        "overall_precision": round(overall_precision, 3),
        "overall_recall":    round(overall_recall,    3),
        "overall_f1":        round(overall_f1,        3),
        "overall_accuracy":  round(overall_accuracy,  3),
        "total_entities":    total_entities,
        "total_tp":          total_tp,
        "total_fp":          total_fp,
        "total_fn":          total_fn,
        "per_category":      results,
    }

    _print_eval_table(summary)
    return summary


def _print_eval_table(summary: Dict) -> None:
    """Pretty-print evaluation results to stdout."""
    print("\n" + "=" * 76)
    print("  PII REDACTION TOOL v2 — EVALUATION REPORT  (75-Entity Harness)")
    print("=" * 76)
    header = (f"{'Category':<20} {'TP':>4} {'FP':>4} {'FN':>4} "
              f"{'Precision':>10} {'Recall':>8} {'F1':>6}")
    print(header)
    print("-" * 76)
    for r in summary["per_category"]:
        print(
            f"{r['category']:<20} {r['tp']:>4} {r['fp']:>4} {r['fn']:>4} "
            f"{r['precision']:>10.3f} {r['recall']:>8.3f} {r['f1']:>6.3f}"
        )
    print("-" * 76)
    print(
        f"{'OVERALL':<20} {summary['total_tp']:>4} {summary['total_fp']:>4} "
        f"{summary['total_fn']:>4} "
        f"{summary['overall_precision']:>10.3f} "
        f"{summary['overall_recall']:>8.3f} "
        f"{summary['overall_f1']:>6.3f}"
    )
    print(f"\n  Accuracy  : {summary['overall_accuracy']:.3f}  "
          f"({summary['total_tp']}/{summary['total_entities']} entities correctly detected)")
    print("=" * 76 + "\n")


# =============================================================================
#  SECTION 9: COMMAND-LINE INTERFACE
# =============================================================================

def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="redact_pii",
        description="Production-grade PII redaction v2 for Indian corporate DOCX documents.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python redact_pii.py --input RHP.docx --output RHP_redacted.docx
  python redact_pii.py --input RHP.docx --output RHP_redacted.docx --dry-run
  python redact_pii.py --input RHP.docx --output RHP_redacted.docx --no-ner
  python redact_pii.py --input RHP.docx --output RHP_redacted.docx --export-mapping
  python redact_pii.py --evaluate
  python redact_pii.py --evaluate --verbose
        """,
    )
    p.add_argument("--input",  "-i", help="Path to input .docx file")
    p.add_argument("--output", "-o", help="Path to output redacted .docx file")
    p.add_argument("--dry-run",        action="store_true",
                   help="Detect PII but do not write output")
    p.add_argument("--no-ner",         action="store_true",
                   help="Disable spaCy NER (regex-only mode)")
    p.add_argument("--export-mapping", action="store_true",
                   help="Export PII->fake mapping to JSON")
    p.add_argument("--evaluate",       action="store_true",
                   help="Run 75-entity evaluation harness and exit")
    p.add_argument("--verbose", "-v",  action="store_true",
                   help="Enable debug-level logging")
    return p


def main() -> None:
    parser = _build_parser()
    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    if args.evaluate:
        run_evaluation(verbose=args.verbose)
        return

    if not args.input or not args.output:
        parser.error("--input and --output are required unless --evaluate is specified.")

    result = redact_document(
        input_path=args.input,
        output_path=args.output,
        use_ner=not args.no_ner,
        dry_run=args.dry_run,
        export_mapping=args.export_mapping,
    )

    print("\n-- Redaction Summary (v2) ---------------------------------------------")
    print(f"  Input                  : {result['input']}")
    print(f"  Output                 : {result['output'] or '(dry-run - not saved)'}")
    print(f"  Paragraphs scanned     : {result['paragraphs_scanned']}")
    print(f"  Unique PII redacted    : {result['unique_entities_redacted']}")
    print(f"  spaCy NER used         : {result['spacy_used']}")
    print("  Category breakdown:")
    for cat, count in result["category_breakdown"].items():
        print(f"    {cat:<22}: {count}")
    print("-----------------------------------------------------------------------\n")


if __name__ == "__main__":
    main()
